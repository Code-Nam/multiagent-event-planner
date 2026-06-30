"""
generate-docx.py — AGEVP Phase 2

Reads a JSON spec file (docx_report shape) and produces a .docx document
with a Heading 1 title, Heading 2 section headings, and Normal paragraphs
for body content.

Usage:
    python scripts/generate-docx.py [--input doc-content/report.json] [--output output/]
"""

import sys
import json
import argparse
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt
except ImportError:
    print("Missing dependency. Run: pip install -r scripts/requirements.txt", file=sys.stderr)
    sys.exit(1)


def slugify(title: str) -> str:
    """Convert a title string to a lowercase underscore slug."""
    return title.lower().replace(" ", "_")


def generate_docx(spec_path: Path, output_dir: Path) -> Path:
    """
    Build a .docx document from a docx_report JSON spec.

    Args:
        spec_path:  Path to the JSON spec file.
        output_dir: Directory where the .docx file will be written.

    Returns:
        Absolute path of the created .docx file.

    Raises:
        SystemExit: on any I/O or validation error.
    """
    try:
        raw = spec_path.read_text(encoding="utf-8")
    except OSError as exc:
        print(f"ERROR: cannot read {spec_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    try:
        spec = json.loads(raw)
    except json.JSONDecodeError as exc:
        print(f"ERROR: invalid JSON in {spec_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    if spec.get("type") != "docx_report":
        print(
            f"ERROR: expected type 'docx_report', got '{spec.get('type')}'",
            file=sys.stderr,
        )
        sys.exit(1)

    title: str = spec.get("title", "Document")
    sections: list = spec.get("sections", [])

    doc = Document()

    doc.add_heading(title, level=1)

    for section in sections:
        heading: str = section.get("heading", "")
        content: str = section.get("content", "")

        if heading:
            doc.add_heading(heading, level=2)
        if content:
            doc.add_paragraph(content)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / f"{slugify(title)}.docx"

    try:
        doc.save(out_path)
    except OSError as exc:
        print(f"ERROR: cannot write {out_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    return out_path.resolve()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a .docx report from a docx_report JSON spec."
    )
    parser.add_argument(
        "--input",
        default="doc-content/report.json",
        help="Path to the JSON spec file (default: doc-content/report.json)",
    )
    parser.add_argument(
        "--output",
        default="output",
        help="Output directory (default: output/)",
    )
    args = parser.parse_args()

    result = generate_docx(
        spec_path=Path(args.input),
        output_dir=Path(args.output),
    )
    print(result)
    sys.exit(0)
