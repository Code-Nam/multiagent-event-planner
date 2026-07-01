"""
generate_docx.py — AGEVP Phase 2

Reads a JSON spec file (docx_report shape) and produces a .docx document
by cloning the AGEVP template structure and replacing placeholder text
with spec content, preserving all direct run-level formatting.

Template body structure (Modele_AGEVP.docx):
  [0]  TABLE  mode d'emploi (removed)
  [1]  p  '' (empty)
  [2]  p  'A G E V P' (branding — kept)
  [3]  p  '[ Titre du document ]' → replaced with spec title
  [4]  p  '[ Sous-titre ]' → replaced with event name
  [5]  p  '' (empty)
  [6]  p  '1  [ Titre de section ]' → cloned per section (gold + blue runs)
  [7]  p  '[ Paragraphe d\'intro ]' → cloned per section (italic gray)
  [32] p  footer line (kept)
  [33] sectPr

Usage:
    python scripts/generate_docx.py <json-spec-path> [--template <path>] [--output <dir>]
"""

import sys
import json
import copy
import argparse
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("Missing dependency. Run: pip install -r scripts/requirements.txt", file=sys.stderr)
    sys.exit(1)

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"
DEFAULT_TEMPLATE = TEMPLATES_DIR / "Modele_AGEVP.docx"

_W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


def _normalize_style_names(doc: Document) -> None:
    # python-docx 1.2+ BabelFish translates specific UI names (e.g. "Heading 1"
    # → "heading 1") before XML lookup. Templates saved by Word/LibreOffice may
    # store the capitalised form. Fix only the aliased names so lookup succeeds.
    from docx.styles.styles import BabelFish
    ui_to_internal = dict(BabelFish.style_aliases)
    for style_elem in doc.styles._element:
        name_elem = style_elem.find(f"{{{_W}}}name")
        if name_elem is not None:
            val = name_elem.get(f"{{{_W}}}val", "")
            if val in ui_to_internal:
                name_elem.set(f"{{{_W}}}val", ui_to_internal[val])


def _replace_first_run_text(para_elem, new_text: str) -> None:
    """Set the text of the first w:t in a paragraph, preserving run formatting."""
    for r in para_elem.iter(f"{{{_W}}}r"):
        t = r.find(f"{{{_W}}}t")
        if t is not None:
            t.text = new_text
            return


def _make_section_heading(heading_tmpl, number: int, title: str):
    """Clone heading template, set number run and title run."""
    cloned = copy.deepcopy(heading_tmpl)
    runs = cloned.findall(f"{{{_W}}}r")
    if len(runs) >= 1:
        t = runs[0].find(f"{{{_W}}}t")
        if t is not None:
            t.text = f"{number}  "
    if len(runs) >= 2:
        t = runs[1].find(f"{{{_W}}}t")
        if t is not None:
            t.text = title
    return cloned


def _make_content_para(content_tmpl, text: str):
    """Clone content template paragraph with new text."""
    cloned = copy.deepcopy(content_tmpl)
    _replace_first_run_text(cloned, text)
    return cloned


def _load_template(template: Path):
    """
    Load AGEVP docx template, extract structural elements, and return
    (doc_with_cover_only, heading_template_elem, content_template_elem, footer_elem).
    """
    doc = Document(str(template))
    _normalize_style_names(doc)

    body = doc.element.body
    orig_elements = list(body)
    n = len(orig_elements)

    # Extract reference elements as deep copies before rebuilding body
    heading_tmpl = copy.deepcopy(orig_elements[6]) if n > 6 else None
    content_tmpl = copy.deepcopy(orig_elements[7]) if n > 7 else None
    footer_tmpl = copy.deepcopy(orig_elements[32]) if n > 32 else None

    # sectPr is always the last child
    sect_pr = body.find(f"{{{_W}}}sectPr")

    # Rebuild body: cover section only (orig[1..5], skip orig[0] = mode d'emploi table)
    for child in list(body):
        body.remove(child)

    cover_indices = range(1, min(6, n))
    for i in cover_indices:
        tag = orig_elements[i].tag.split("}")[1]
        if tag != "sectPr":
            body.append(copy.deepcopy(orig_elements[i]))

    if sect_pr is not None:
        body.append(sect_pr)

    return doc, heading_tmpl, content_tmpl, footer_tmpl


def generate_docx(spec_path: Path, output_dir: Path, template: Path | None = None) -> Path:
    """
    Build a .docx document from a docx_report JSON spec.

    Args:
        spec_path:  Path to the JSON spec file.
        output_dir: Directory where the .docx file will be written.
        template:   Optional path to a .docx template for branding.

    Returns:
        Absolute path of the created .docx file.
    """
    try:
        raw = spec_path.read_text(encoding="utf-8")
    except OSError as exc:
        print(f"ERROR: cannot read {spec_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    try:
        spec = json.loads(raw)
    except json.JSONDecodeError as exc:
        print(f"ERROR: invalid JSON in {spec_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    if spec.get("type") != "docx_report":
        print(
            f"ERROR: expected type 'docx_report', got '{spec.get('type')}'",
            file=sys.stderr,
        )
        sys.exit(1)

    title: str = spec.get("title", "Document")
    sections: list = spec.get("sections", [])
    event: str = spec.get("event", "")
    subtitle: str = spec.get("subtitle", event)

    resolved_template = template if template is not None else DEFAULT_TEMPLATE

    if resolved_template and resolved_template.exists():
        doc, heading_tmpl, content_tmpl, footer_tmpl = _load_template(resolved_template)

        body = doc.element.body
        sect_pr = body.find(f"{{{_W}}}sectPr")
        if sect_pr is not None:
            body.remove(sect_pr)

        # Cover: body[0]=empty, [1]=A G E V P, [2]=title placeholder, [3]=subtitle, [4]=empty
        cover_children = [c for c in body if c.tag.split("}")[1] != "sectPr"]
        if len(cover_children) > 2:
            _replace_first_run_text(cover_children[2], title)
        if len(cover_children) > 3:
            _replace_first_run_text(cover_children[3], subtitle)

        # Append section blocks
        for i, section in enumerate(sections, 1):
            heading_text: str = section.get("heading", "")
            content_text: str = section.get("content", "")

            if heading_tmpl is not None:
                body.append(_make_section_heading(heading_tmpl, i, heading_text))
            if content_tmpl is not None:
                body.append(_make_content_para(content_tmpl, content_text))

        # Append footer then sectPr
        if footer_tmpl is not None:
            body.append(copy.deepcopy(footer_tmpl))
        if sect_pr is not None:
            body.append(sect_pr)

    else:
        # Fallback: plain document without template branding
        doc = Document()
        doc.add_heading(title, level=1)
        for section in sections:
            heading_text = section.get("heading", "")
            content_text = section.get("content", "")
            if heading_text:
                doc.add_heading(heading_text, level=2)
            if content_text:
                doc.add_paragraph(content_text)

    output_dir.mkdir(parents=True, exist_ok=True)
    out_path = output_dir / f"{spec_path.stem}.docx"

    try:
        doc.save(out_path)
    except OSError as exc:
        print(f"ERROR: cannot write {out_path}: {exc}", file=sys.stderr)
        sys.exit(1)

    return out_path.resolve()


if __name__ == "__main__":
    parser = argparse.ArgumentParser(
        description="Generate a .docx report from a docx_report JSON spec."
    )
    parser.add_argument("spec", help="Path to the JSON spec file")
    parser.add_argument(
        "--template",
        default=None,
        help=f"Path to a .docx template for branding (default: {DEFAULT_TEMPLATE})",
    )
    parser.add_argument(
        "--output",
        default="output",
        help="Output directory (default: output/)",
    )
    args = parser.parse_args()

    result = generate_docx(
        spec_path=Path(args.spec),
        output_dir=Path(args.output),
        template=Path(args.template) if args.template else None,
    )
    print(result)
    sys.exit(0)
